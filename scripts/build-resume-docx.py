#!/usr/bin/env python3
"""Build an ATS-optimised combined DOCX resume from resume.json.

Content + ordering come from resume_blocks.build_blocks() (shared with the PDF
builder) so a single source of truth produces every output format. This module
owns only the DOCX rendering of those blocks.

ATS-optimisation choices (research from Jobscan, Greenhouse parser guide, Refhub AU,
2026 — see docs/resume-generation.md):

- Single-column layout. No tables, text boxes, columns, headers, footers, images, icons.
- Standard fonts only: Calibri 11pt body, 14pt section headings, 20pt name.
- Standard section names ATS look for: "Professional Summary", "Work Experience",
  "Education", "Certifications", "Skills".
- Reverse-chronological work history. Dates "Month YYYY" (long form, unambiguous).
- Contact info inline at the top with text labels (Phone:/Email:/LinkedIn:), no icons.
- ASCII bullet `* ` for highest compatibility (some ATS misread Unicode bullets).
- A4 page (AU standard) with 2.5 cm margins. python-docx defaults to US Letter, so set explicitly.
- Document proofing language set to Australian English (en-AU).

Run: python3 scripts/build-resume-docx.py [--real-contact]

By default `basics.email` renders the obfuscated form (`cv [at] paulharvey [dot] com
[dot] au`) and the phone is omitted — sufficient for spam-scraper resistance, decodable
by humans + modern ATS. For direct ATS submissions needing an unambiguous contact field,
run with `--real-contact`: output goes to `Paul Harvey - Resume (full).docx` (NOT committed).
"""

from __future__ import annotations

import argparse
import datetime as dt
import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Shared helpers + the format-agnostic block model. scripts/ is normally sys.path[0]
# when run as `python3 scripts/build-resume-docx.py`; insert explicitly to be robust.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from resume_blocks import build_blocks, to_ascii  # noqa: E402


REPO = Path(__file__).resolve().parent.parent
PUBLIC = REPO / "public"               # wrangler deploy root — everything here is published
RESUME_JSON = PUBLIC / "resume.json"
OUT_DOCX = PUBLIC / "cv" / "docs" / "Paul Harvey - Resume.docx"

# Fonts: Calibri is the most widely-tested font for ATS parsing.
FONT_BODY = "Calibri"
FONT_HEADING = "Calibri"

# ASCII bullet for max compatibility. Some ATS still trip on `•`.
BULLET = "* "


def set_run(run, *, size=11, bold=False, italic=False, font=FONT_BODY):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    # Set East-Asian font too so Word doesn't substitute on non-Latin Word installs
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:cs"), font)


def add_para(doc, text="", *, size=11, bold=False, italic=False, font=FONT_BODY,
             space_before=0, space_after=2, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(to_ascii(text))
        set_run(r, size=size, bold=bold, italic=italic, font=font)
    return p


def add_heading(doc, text):
    """Section heading. Capitalised standard ATS names. Bold + larger size only."""
    return add_para(doc, text, size=14, bold=True, font=FONT_HEADING,
                    space_before=8, space_after=3)


def add_subheading(doc, text):
    """Role/sub-section heading."""
    return add_para(doc, text, size=12, bold=True, space_before=6, space_after=1)


def add_bullet(doc, text, indent_level=0):
    """Plain bulleted line. Uses ASCII '* ' prefix in text rather than Word's list-style,
    because ATS often strip Word list formatting but preserve in-paragraph text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.2 + indent_level * 0.2)
    r = p.add_run(BULLET + to_ascii(text))
    set_run(r, size=11)
    return p


def add_hr(doc):
    """Horizontal rule via single paragraph border. ATS-safe."""
    p = add_para(doc, "", size=4, space_before=2, space_after=4)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_document_language(doc, lang="en-AU"):
    """Set proofing language so Word spell-checks Australian English.

    Sets it on both docDefaults (inherited by all text) and the Normal style.
    Without this, python-docx documents default to en-US and Word flags
    AU spellings (organise, recognised, behaviour) as errors.
    """
    styles = doc.styles.element  # <w:styles>

    docDefaults = styles.find(qn("w:docDefaults"))
    if docDefaults is None:
        docDefaults = OxmlElement("w:docDefaults")
        styles.insert(0, docDefaults)
    rPrDefault = docDefaults.find(qn("w:rPrDefault"))
    if rPrDefault is None:
        rPrDefault = OxmlElement("w:rPrDefault")
        docDefaults.append(rPrDefault)
    rPr = rPrDefault.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPrDefault.append(rPr)
    for el in (rPr, doc.styles["Normal"].element.get_or_add_rPr()):
        langEl = el.find(qn("w:lang"))
        if langEl is None:
            langEl = OxmlElement("w:lang")
            el.append(langEl)
        langEl.set(qn("w:val"), lang)


def set_core_metadata(doc, *, name, label, last_modified, title=None):
    """Author the document properties as a person, not the toolchain.

    python-docx ships placeholder core properties — creator "python-docx",
    description "generated by python-docx", and a frozen 2013 created/modified
    date. Those self-identify the file as script-generated to any Properties
    pane or metadata scanner. Overwrite with truthful, person-authored values.
    Producer-side honesty: these are all real (Paul authored the content); we
    only strip the toolchain fingerprint, we don't spoof an editor.

    `title` defaults to "<name> — Resume"; pass it for other document kinds
    (e.g. a cover letter).
    """
    cp = doc.core_properties
    cp.author = name
    cp.last_modified_by = name
    cp.title = title or f"{name} - Resume"
    cp.subject = label
    cp.comments = ""  # clears "generated by python-docx"
    try:
        created = dt.datetime.strptime(last_modified, "%Y-%m-%d")
    except (TypeError, ValueError):
        created = None
    if created is not None:
        cp.created = created
        cp.modified = created


def normalize_docx(path, fixed=(2020, 1, 1, 0, 0, 0)):
    """Rewrite the .docx zip with fixed entry timestamps so identical content produces
    identical bytes. python-docx stamps each zip entry with the build time, which would
    otherwise churn the committed artifact on every rebuild. Entry order is preserved."""
    tmp = path.with_name(path.name + ".tmp")
    with zipfile.ZipFile(path) as zin, \
         zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            zi = zipfile.ZipInfo(item.filename, date_time=fixed)
            zi.compress_type = zipfile.ZIP_DEFLATED
            zi.external_attr = item.external_attr
            zout.writestr(zi, zin.read(item.filename))
    tmp.replace(path)


def render_block(doc, blk):
    """Map one format-agnostic block to its DOCX paragraph(s)."""
    t = blk["t"]
    if t == "heading":
        add_heading(doc, blk["text"])
    elif t == "sub":
        add_subheading(doc, blk["text"])
    elif t == "para":
        add_para(doc, blk["text"], size=blk["size"], bold=blk["bold"],
                 italic=blk["italic"], space_before=blk["sb"], space_after=blk["sa"])
    elif t == "bullet":
        add_bullet(doc, blk["text"], indent_level=blk["level"])
    elif t == "hr":
        add_hr(doc)
    elif t == "spacer":
        add_para(doc, "", size=blk["size"])
    else:
        raise ValueError(f"Unknown block type: {t!r}")


def main():
    parser = argparse.ArgumentParser(description="Build ATS-friendly DOCX from resume.json")
    parser.add_argument("--real-contact", action="store_true",
                        help="Use the real email + phone (for direct ATS submissions). "
                             "Output goes to 'Paul Harvey - Resume (full).docx', NOT the public file.")
    args = parser.parse_args()

    data = json.loads(RESUME_JSON.read_text())
    basics = data["basics"]

    # Output path depends on contact mode. Contact resolution itself (email/phone, which
    # live in resume_blocks, not resume.json) happens inside build_blocks().
    if args.real_contact:
        # Private real-contact variant: write OUTSIDE public/ so wrangler never publishes it.
        out_path = REPO / "out" / "full-contact" / "Paul Harvey - Resume (full).docx"
    else:
        out_path = OUT_DOCX

    doc = Document()

    # A4 page (AU standard) — python-docx defaults to US Letter. Margins 2.5 cm all round.
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        # Explicitly disable headers/footers — many ATS skip them
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True

    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    set_document_language(doc, "en-AU")

    # Render the shared block model.
    for blk in build_blocks(data, real_contact=args.real_contact):
        render_block(doc, blk)

    # ====== FOOTER NOTE: canonical link (build timestamp lives here, not in build_blocks) ======
    add_para(doc, "", size=8)
    add_para(doc,
             f"Live interactive resume: {basics['url']} | Updated {data['meta']['lastModified']} | "
             f"Source: resume.json v{data['meta']['version']}",
             size=9, italic=True)

    set_core_metadata(doc, name=basics["name"],
                       label=basics.get("label", ""),
                       last_modified=data["meta"]["lastModified"])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    normalize_docx(out_path)  # deterministic bytes — no churn on rebuild

    work = data.get("work", [])
    recent = sum(1 for w in work if not w.get("earlierCareer"))
    earlier = len(work) - recent
    print(f"Wrote {out_path}")
    print(f"  Size: {out_path.stat().st_size:,} bytes")
    print(f"  Roles (recent + earlier): {recent} + {earlier}")
    print(f"  Certifications: {len(data.get('certifications', []))}")
    print(f"  Contact mode: {'REAL' if args.real_contact else 'obfuscated'}")


if __name__ == "__main__":
    main()
