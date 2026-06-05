"""Tests for build-resume-docx.py: DOCX rendering helpers + main()."""

from __future__ import annotations

import zipfile

import pytest
from docx import Document
from docx.oxml.ns import qn


def test_set_run_creates_and_reuses_rfonts(docx_mod):
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("hello")
    docx_mod.set_run(r, size=13, bold=True, italic=True, font="Calibri")
    rfonts = r._element.get_or_add_rPr().find(qn("w:rFonts"))
    assert rfonts.get(qn("w:ascii")) == "Calibri"
    assert r.font.bold is True and r.font.italic is True
    # second call hits the "rfonts is not None" branch
    docx_mod.set_run(r, size=11, font="Arial")
    rfonts2 = r._element.get_or_add_rPr().find(qn("w:rFonts"))
    assert rfonts2.get(qn("w:ascii")) == "Arial"


def test_add_para_variants(docx_mod):
    doc = Document()
    p = docx_mod.add_para(doc, "text", size=12, bold=True, italic=True,
                          space_before=2, space_after=3)
    assert p.runs[0].text == "text"
    # empty text -> no run
    empty = docx_mod.add_para(doc, "")
    assert empty.runs == []
    # explicit alignment branch
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    aligned = docx_mod.add_para(doc, "c", align=WD_ALIGN_PARAGRAPH.CENTER)
    assert aligned.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_add_heading_subheading_bullet(docx_mod):
    doc = Document()
    h = docx_mod.add_heading(doc, "Section")
    assert h.runs[0].text == "Section"
    s = docx_mod.add_subheading(doc, "Role")
    assert s.runs[0].text == "Role"
    b = docx_mod.add_bullet(doc, "point", indent_level=0)
    assert b.runs[0].text.startswith("* ")
    b1 = docx_mod.add_bullet(doc, "subpoint", indent_level=1)
    assert b1.paragraph_format.left_indent > b.paragraph_format.left_indent


def test_add_hr(docx_mod):
    doc = Document()
    p = docx_mod.add_hr(doc)
    pBdr = p._element.get_or_add_pPr().find(qn("w:pBdr"))
    assert pBdr is not None


def test_set_document_language_idempotent(docx_mod):
    doc = Document()
    docx_mod.set_document_language(doc, "en-AU")
    # second call exercises the "already exists" branches
    docx_mod.set_document_language(doc, "en-GB")
    styles = doc.styles.element
    docDefaults = styles.find(qn("w:docDefaults"))
    rPr = docDefaults.find(qn("w:rPrDefault")).find(qn("w:rPr"))
    assert rPr.find(qn("w:lang")).get(qn("w:val")) == "en-GB"


def test_set_core_metadata_valid_date(docx_mod):
    doc = Document()
    docx_mod.set_core_metadata(doc, name="Jane", label="Dev",
                               last_modified="2026-06-04")
    cp = doc.core_properties
    assert cp.author == "Jane"
    assert cp.last_modified_by == "Jane"
    assert cp.title == "Jane - Resume"
    assert cp.subject == "Dev"
    assert cp.created.year == 2026 and cp.created.month == 6


def test_set_core_metadata_custom_title(docx_mod):
    doc = Document()
    docx_mod.set_core_metadata(doc, name="Jane", label="Dev",
                               last_modified="2026-06-04", title="Cover Letter")
    assert doc.core_properties.title == "Cover Letter"


def test_set_core_metadata_bad_date_skips_created(docx_mod):
    doc = Document()
    # invalid string -> ValueError -> created stays None (not set)
    docx_mod.set_core_metadata(doc, name="Jane", label="Dev", last_modified="not-a-date")
    # None -> TypeError -> also skipped
    docx_mod.set_core_metadata(doc, name="Jane", label="Dev", last_modified=None)
    assert doc.core_properties.author == "Jane"


def test_render_block_all_types(docx_mod, blocks_mod):
    doc = Document()
    for blk in [
        blocks_mod._h("H"),
        blocks_mod._sub("S"),
        blocks_mod._p("P", size=11),
        blocks_mod._b("B"),
        blocks_mod._hr(),
        blocks_mod._sp(6),
    ]:
        docx_mod.render_block(doc, blk)
    # at least one paragraph rendered per block
    assert len(doc.paragraphs) >= 6


def test_render_block_unknown_raises(docx_mod):
    doc = Document()
    with pytest.raises(ValueError, match="Unknown block type"):
        docx_mod.render_block(doc, {"t": "bogus"})


def test_normalize_docx_deterministic(docx_mod, tmp_path, blocks_mod):
    path = tmp_path / "r.docx"
    doc = Document()
    docx_mod.render_block(doc, blocks_mod._p("hello", size=11))
    doc.save(path)
    docx_mod.normalize_docx(path)
    first = path.read_bytes()
    # all zip entries pinned to the fixed timestamp
    with zipfile.ZipFile(path) as z:
        assert all(info.date_time == (2020, 1, 1, 0, 0, 0) for info in z.infolist())
    # re-normalizing yields identical bytes
    docx_mod.normalize_docx(path)
    assert path.read_bytes() == first


def test_main_default(docx_mod, monkeypatch, tmp_path):
    out = tmp_path / "Resume.docx"
    monkeypatch.setattr(docx_mod, "OUT_DOCX", out)
    monkeypatch.setattr("sys.argv", ["build-resume-docx.py"])
    docx_mod.main()
    assert out.exists() and out.stat().st_size > 0
    # output is a valid zip (docx)
    assert zipfile.is_zipfile(out)


def test_main_real_contact(docx_mod, monkeypatch, tmp_path, capsys):
    monkeypatch.setattr(docx_mod, "REPO", tmp_path)
    monkeypatch.setattr("sys.argv", ["build-resume-docx.py", "--real-contact"])
    docx_mod.main()
    out = tmp_path / "out" / "full-contact" / "Paul Harvey - Resume (full).docx"
    assert out.exists()
    assert "REAL" in capsys.readouterr().out
